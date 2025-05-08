import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

# -------- Gráfica 1 --------
x1 = [1, 2, 3, 4, 5, 6]
y1 = [8.65, 17.30, 25.95, 34.63, 43.31, 51.95]

plt.plot(x1, y1, marker='o')
plt.xticks(range(1, 8))
plt.yticks(range(10, 70, 10))
plt.xlabel('m [g]')
plt.ylabel('H [cm]')
plt.grid(True)
plt.savefig('imagenes/grafica1.png', dpi=300)
plt.close()

# -------- Gráfica 2 --------
x2 = [1, 2, 3, 4, 5, 6]
y2 = [1.22, 4.90, 10.40, 19.52, 30.71, 43.75]

plt.plot(x2, y2, marker='o')
plt.xticks(range(1, 8))
plt.yticks(range(10, 70, 10))
plt.xlabel('D [cm]')
plt.ylabel('m [g]')
plt.grid(True)
plt.savefig('imagenes/grafica2.png', dpi=300)
plt.close()

# -------- Gráfica 3 --------
x3 = [val ** 2 for val in x2]
y3 = y2

plt.plot(x3, y3, marker='o')
plt.xticks(range(0, 70, 10))
plt.yticks(range(10, 70, 10))
plt.xlabel('D² [cm²]')
plt.ylabel('m [g]')
plt.grid(True)
plt.savefig('imagenes/grafica3.png', dpi=300)
plt.close()

# -------- Gráfica 4 --------
x4 = np.log10(x2)
y4 = np.log10(y2)

plt.plot(x4, y4, marker='o')
plt.xticks(np.arange(0, 1.1, 0.1))
plt.yticks(np.arange(0, 1.8, 0.1))
plt.xlabel('x = log10(D [cm])')
plt.ylabel('y = log10(m [g])')
plt.grid(True)
plt.savefig('imagenes/grafica4.png', dpi=300)
plt.close()

# -------- Gráfica 5 --------
x5 = [0.713, 0.998, 1.501, 1.746, 1.905, 2.222]
y5 = [1.47, 4.50, 13.75, 21.70, 28.20, 44.75]

plt.plot(x5, y5, marker='o')
plt.xticks([round(i * 0.1, 1) for i in range(0, 26)])  # 0.0 to 2.5 by 0.1
plt.yticks(range(10, 60, 10))  # 10 to 50
plt.xlabel('D [cm]')
plt.ylabel('m [g]')
plt.grid(True)
plt.savefig('imagenes/grafica5.png', dpi=300)
plt.close()

# -------- Gráfica 6 --------
x6 = [round(val**2, 4) for val in x5]  # Z = D² [cm²]
y6 = y5  # m [g]

plt.plot(x6, y6, marker='o')
plt.yticks(range(10, 70, 10))  # 10 a 60
plt.xticks(range(1, 7))        # 1 a 6
plt.xlabel('Z [cm²]')
plt.ylabel('m [g]')
plt.grid(True)
plt.savefig('imagenes/grafica6.png', dpi=300)
plt.close()

# -------- Gráfica 7 --------
x7 = [round(val**3, 4) for val in x5]  # W = D³ [cm³]
y7 = y5  # m [g]

plt.plot(x7, y7, marker='o')
plt.xticks(range(1, 13))  # 1 a 12
plt.yticks(range(10, 80, 10))  # 10 a 70
plt.xlabel('W [cm³]')
plt.ylabel('m [g]')
plt.grid(True)
plt.savefig('imagenes/grafica7.png', dpi=300)
plt.close()

# -------- Gráfica 8 --------
x8 = np.log10(x5)  # log10(D [cm])
y8 = np.log10(y5)  # log10(m [g])

plt.plot(x8, y8, marker='o')
plt.xlabel('x = log10(D [cm])')
plt.ylabel('y = log10(m [g])')
plt.xticks(np.arange(-1, 2.5, 0.5))
plt.yticks(np.arange(-1, 3.5, 0.5))
plt.grid(True)
plt.savefig('imagenes/grafica8.png', dpi=300)
plt.close()


# -------- Insertar en DOCX --------
doc = Document('plantilla.docx')

for para in doc.paragraphs:
    if '<<GRAFICA1>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica1.png', width=Inches(4))
    if '<<GRAFICA2>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica2.png', width=Inches(4))
    if '<<GRAFICA3>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica3.png', width=Inches(4))
    if '<<GRAFICA4>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica4.png', width=Inches(4))
    if '<<GRAFICA5>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica5.png', width=Inches(4))
    if '<<GRAFICA6>>' in para.text:     
        para.clear()
        para.add_run().add_picture('imagenes/grafica6.png', width=Inches(4))
    if '<<GRAFICA7>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica7.png', width=Inches(4))
    if '<<GRAFICA8>>' in para.text:
        para.clear()
        para.add_run().add_picture('imagenes/grafica8.png', width=Inches(4))

doc.save('informePythonFrancoPrieto.docx')
